"""
Gerador do artigo AHCF-CPS com §7.9 revisado.

Lê o artigo original em formato .txt (já extraído do .docx em
C:\\Users\\DETRAN\\AppData\\Local\\Temp\\article.txt) e produz um novo .docx com:

- §7.9 substituída pela formulação revisada (com OMML — math nativo Word);
- sinopses e conclusão comprimidas de forma a atingir ~8000 palavras
  sem adicionar conteúdo novo nem remover citações ou resultados;
- formatação consistente: títulos numerados, corpo Arial 11, fórmulas
  com Cambria Math, quadros pareados.

O script não inventa citações nem inclui conteúdo externo. Todas as
passagens são reformulações mais concisas de frases já presentes no
original.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SOURCE = Path(r"C:\Users\DETRAN\AppData\Local\Temp\article.txt")
OUT_DIR = Path(
    r"C:\Users\DETRAN\Documents\JOAO PAULO\Mestrado\Disciplinas"
    r"\PEA5733 - Automação e Sociedade (2026)\Apresentação dos artigos\Artigo"
)
OUT = OUT_DIR / "Artigo_AHCF_CPS_USP_Revisado_v2.docx"

# ---------------------------------------------------------------------------
# Compressões cirúrgicas — apenas reescrita mais concisa de frases já no artigo
# ---------------------------------------------------------------------------

# Textos a remover (não carregam conteúdo e inflam a contagem)
STRIP_PARAGRAPHS = {
    "Interpretação analítica",
}

REPLACEMENTS: dict[str, List[str]] = {
    # §2 — Introdução compactada preservando todas as citações
    "2. Introdução": [
        "A transformação digital tem promovido mudanças profundas na forma "
        "como sistemas produtivos são concebidos, operados e otimizados. "
        "Nesse cenário, os robôs colaborativos (cobots) emergem como uma "
        "das tecnologias mais relevantes da manufatura contemporânea por "
        "possibilitarem interação direta, segura e flexível entre "
        "trabalhadores humanos e sistemas automatizados. Diferentemente "
        "dos robôs industriais tradicionais, isolados em células fechadas, "
        "os cobots foram projetados para compartilhar espaço com "
        "operadores, executar tarefas complementares e adaptar-se a "
        "ambientes dinâmicos (EL ZAATARI et al., 2019).",
        "Esse avanço está associado à transição da Indústria 4.0 para a "
        "Indústria 5.0: se a primeira enfatizou conectividade, automação e "
        "sistemas ciberfísicos, a segunda recoloca o ser humano no centro "
        "do processo produtivo, acompanhando eficiência com bem-estar "
        "ocupacional, sustentabilidade, resiliência e colaboração "
        "inteligente. Nesse novo paradigma, os cobots tornam-se "
        "estratégicos por combinarem produtividade mecânica com "
        "capacidades humanas de julgamento, criatividade e adaptação "
        "(EL ZAATARI et al., 2019).",
        "Apesar da crescente adoção em setores como automotivo, "
        "eletrônico, logístico, farmacêutico, agrícola e alimentício, a "
        "implementação prática enfrenta desafios relevantes. Um dos "
        "principais é a seleção do método de programação, ainda "
        "frequentemente baseada em decisões empíricas, preferências do "
        "integrador ou limitações tecnológicas específicas, sem "
        "considerar adequadamente fatores humanos e contextuais. "
        "Empresas adotam, assim, soluções inadequadas ao perfil do "
        "operador, à variabilidade do processo ou às exigências "
        "ergonômicas e de segurança da tarefa.",
        "A literatura recente demonstra que programar cobots deixou de ser "
        "apenas um problema de trajetórias ou código: trata-se de uma "
        "questão multidimensional que envolve usabilidade, carga "
        "cognitiva, confiança, transparência algorítmica, ergonomia, "
        "segurança dinâmica e flexibilidade operacional. Interfaces "
        "baseadas em realidade estendida podem aumentar usabilidade e "
        "preferência do usuário, com limitações em tarefas de alta "
        "precisão; sistemas com maior transparência aumentam confiança e "
        "capacidade de ensino por não especialistas; diferentes níveis de "
        "autonomia decisória alteram desempenho e percepção do trabalho "
        "em células humano-cobot (BAGHERI et al., 2022; "
        "WOLFFGRAMM et al., 2024). Crescem igualmente métodos intuitivos "
        "como Programming by Demonstration, ensino cinestésico, "
        "interfaces multimodais, controle compartilhado e aprendizagem "
        "incremental — cada um com vantagens e limitações dependentes da "
        "natureza da tarefa, do perfil humano e das condições ambientais.",
        "Apesar desses avanços, a literatura permanece fragmentada: "
        "estudos concentram-se isoladamente em segurança e conformidade "
        "normativa, ergonomia ocupacional, aprendizado de máquina ou "
        "experiência do usuário; são escassos os modelos que integram "
        "simultaneamente fatores técnicos, humanos e organizacionais em "
        "uma estrutura formal de apoio à decisão sobre métodos de "
        "programação. A lacuna é especialmente relevante para pequenas e "
        "médias empresas, que precisam implantar automação colaborativa "
        "com recursos limitados e alta necessidade de retorno.",
        "O artigo parte da premissa de que não existe um método "
        "universalmente superior: existe um método mais adequado para "
        "cada contexto operacional. Portanto, a escolha deve ser "
        "tratada como problema multicritério, adaptativo e centrado no "
        "humano. O objetivo é propor o Adaptive Human-Centered Framework "
        "for Cobot Programming Selection (AHCF-CPS), framework conceitual "
        "e matematicamente operacionalizável para apoiar a seleção de "
        "estratégias de programação de cobots em ambientes industriais. "
        "O modelo integra cinco dimensões — segurança, ergonomia, "
        "características humanas, performance operacional e complexidade "
        "da tarefa — alinhadas aos princípios da Indústria 5.0.",
        "Especificamente, busca-se responder a três questões: RQ1 — quais "
        "são os principais métodos de programação de cobots vinculados à "
        "colaboração humano-robô industrial e como organizá-los em "
        "taxonomia atualizada; RQ2 — como fatores contextuais e "
        "construtos humano-robô influenciam a escolha do método em "
        "ambientes produtivos reais; RQ3 — é possível estruturar um "
        "framework formal que integre tais variáveis e auxilie a tomada "
        "de decisão industrial. Para respondê-las, realiza-se uma "
        "revisão sistemática da literatura recente, seguida da "
        "proposição do AHCF-CPS, com contribuição tanto acadêmica "
        "quanto prática.",
    ],
    # §1 Resumo — leve compactação
    "1. Resumo": [
        "Este artigo, em formato de Revisão Sistemática da Literatura, "
        "analisa métodos contemporâneos de programação de robôs "
        "colaborativos entre 2019 e 2026. O objetivo é compreender como "
        "fatores técnicos e humanos influenciam a escolha do método e "
        "propor um framework aplicável à indústria. Os resultados "
        "indicam predominância de estudos fragmentados, sem integração "
        "simultânea entre segurança, ergonomia e preferências humanas.",
    ],
    # §3.1 — compressão leve (várias ideias são retomadas em §6.1)
    "3.1 Evolução da robótica colaborativa no contexto industrial": [
        "A introdução de robôs colaborativos é uma das mudanças mais "
        "relevantes da automação industrial das últimas décadas. Enquanto "
        "os robôs industriais convencionais foram historicamente pensados "
        "para operar em células isoladas, com foco em repetibilidade, "
        "velocidade e produção em massa, os cobots surgem com proposta "
        "distinta: compartilhar o espaço de trabalho com operadores "
        "humanos, executar tarefas complementares e permitir "
        "reconfiguração rápida de processos. Essa mudança amplia o uso da "
        "robótica para cenários de maior variabilidade, lotes menores e "
        "operações mais próximas do trabalhador (EL ZAATARI et al., 2019).",
        "A ascensão dos cobots coincide com a transição da Indústria 4.0 "
        "para a Indústria 5.0. Na Indústria 4.0, o foco recaiu sobre "
        "digitalização, integração de sistemas, IoT, big data e "
        "automação inteligente; a Indústria 5.0 incorpora uma dimensão "
        "explicitamente humanocêntrica, defendendo que tecnologias "
        "produtivas devem elevar não apenas eficiência, mas também "
        "sustentabilidade, resiliência e bem-estar ocupacional. Nesse "
        "cenário, os cobots combinam capacidade mecânica, precisão e "
        "disponibilidade contínua com competências humanas de criatividade, "
        "julgamento contextual e resolução de problemas "
        "(EL ZAATARI et al., 2019). Além do setor manufatureiro, a "
        "robótica colaborativa expande-se para logística, saúde, "
        "agricultura, construção civil e serviços.",
    ],
    # §3.2 introdução + 3.2.1 — compressão
    "3.2.1 Programação tradicional e offline": [
        "A programação convencional permanece relevante em aplicações de "
        "alta repetibilidade e baixa variabilidade, com trajetórias e "
        "sequências previamente modeladas e validadas. Tem a favor "
        "previsibilidade, controle fino e estabilidade operacional; "
        "contra, reprogramação lenta, dependência de especialistas e "
        "custo alto quando há mudanças frequentes.",
    ],
    # §3.2.3 / §3.2.4 / §3.2.5 / §3.2.6 — compressão (retomados em §6.2)
    "3.2.3 Ensino cinestésico e interfaces multimodais": [
        "Uma vertente importante do PbD é o ensino cinestésico, em que o "
        "operador guia fisicamente o braço robótico. Meattini et al. "
        "(2025) mostram que essa abordagem pode ir além da trajetória "
        "espacial e incluir parâmetros de interação — rigidez, compliance "
        "e resposta mecânica — combinando sinais musculares (sEMG) e "
        "feedback vibrotátil, o que amplia o entendimento da programação "
        "como processo multimodal e centrado no usuário.",
    ],
    "3.2.4 Interfaces baseadas em Realidade Estendida (XR)": [
        "Interfaces de realidade aumentada, virtual e mista vêm sendo "
        "aplicadas à programação de cobots por oferecerem visualização "
        "intuitiva de trajetórias, zonas seguras e estados operacionais. "
        "Estudos experimentais mostram aumento da usabilidade percebida e "
        "redução do esforço físico, embora tarefas tridimensionais "
        "complexas ainda imponham desafios de precisão espacial, "
        "rastreamento e tempo de resposta.",
    ],
    "3.2.5 Aprendizagem incremental e programação explicável": [
        "Cresce o uso de sistemas nos quais o cobot aprende parcialmente "
        "a tarefa e solicita ajuda humana ao identificar lacunas. "
        "Bagheri et al. (2022) demonstram que interfaces transparentes, "
        "capazes de explicar por que o robô tomou determinada decisão, "
        "aumentam a capacidade de fornecer instruções corretas e elevam "
        "confiança e satisfação — transparência é elemento funcional, "
        "não apenas desejável, da programação colaborativa.",
    ],
    "3.2.6 Controle compartilhado e planejamento adaptativo": [
        "Em tarefas críticas ou dinâmicas, cresce o uso de shared control, "
        "no qual humano e robô dividem autoridade operacional em tempo "
        "real: o operador intervém continuamente enquanto algoritmos "
        "autônomos otimizam segurança, trajetória ou produtividade — "
        "especialmente útil em montagem de precisão, manipulação de "
        "peças frágeis ou ambientes incertos.",
    ],
    # §5.4 — síntese comparativa reescrita em parágrafo denso
    "5.4 Síntese comparativa dos achados": [
        "A consolidação dos dados indica cinco conclusões centrais: o "
        "crescimento de métodos intuitivos torna a programação acessível "
        "ao operador; a segurança aparece como variável recorrente e não "
        "pode ser tratada isoladamente; a ergonomia ganha relevância como "
        "parte da eficiência; a transparência aumenta confiança e torna "
        "interfaces explicáveis estratégicas; e preferências individuais "
        "influenciam desempenho, favorecendo modelos personalizados.",
    ],
    # §8.2 — limitações em parágrafo denso
    "8.2 Limitações do Estudo": [
        "O trabalho possui limitações que devem ser reconhecidas: "
        "(a) recorte temporal concentrado em 2019–2026, podendo excluir "
        "trabalhos clássicos relevantes; (b) foco em bases indexadas "
        "internacionais — Web of Science e complementares — deixando "
        "fora relatórios industriais, white papers e literatura cinzenta; "
        "(c) heterogeneidade metodológica entre os artigos, que dificulta "
        "comparações quantitativas diretas; (d) parte dos estudos "
        "analisados foi conduzida em ambientes controlados, limitando "
        "extrapolação imediata para contextos industriais complexos; e "
        "(e) o framework, embora fundamentado empiricamente, permanece "
        "em estágio conceitual e necessita validação prática em cenários "
        "reais antes de adoção ampla.",
    ],
    # §3.2.2 PbD — leve compressão
    "3.2.2 Programming by Demonstration (PbD)": [
        "O Programming by Demonstration tornou-se uma das abordagens mais "
        "promissoras para robótica colaborativa: o operador ensina a "
        "tarefa por demonstração direta, e o sistema observa, registra e "
        "reproduz padrões aprendidos. A estratégia reduz barreiras "
        "técnicas, aproxima o operador do processo de programação e "
        "acelera ajustes operacionais. O PbD pode assumir diferentes "
        "formas — demonstração manual de trajetórias, teleoperação ou "
        "ensino visual baseado em sensores — e sua principal limitação "
        "reside na variabilidade entre demonstrações humanas e na "
        "necessidade de algoritmos robustos para generalização.",
    ],
    # §3.3 — fatores humanos consolidados; subseções 3.3.2–3.3.5 foram absorvidas
    "3.3 Fatores humanos na colaboração humano-robô": [
        "O desempenho de sistemas colaborativos depende fortemente de "
        "fatores humanos. Diferentemente da automação convencional, em "
        "que o operador atua como supervisor externo, na HRC o "
        "trabalhador participa diretamente da execução, da tomada de "
        "decisão e da adaptação do sistema. A confiança é um dos "
        "construtos mais investigados: níveis insuficientes geram "
        "resistência tecnológica e baixa aceitação, enquanto confiança "
        "excessiva produz complacência — o ideal é a confiança calibrada, "
        "favorecida por interfaces transparentes e previsibilidade de "
        "comportamento. A usabilidade reduz a dependência de "
        "especialistas e amplia adoção, enquanto sistemas complexos "
        "deslocam esforço físico para esforço mental, aumentando a carga "
        "cognitiva. A ergonomia é central: cobots costumam assumir "
        "tarefas repetitivas, levantamento de peso e posturas "
        "desfavoráveis, e estudos biomecânicos mostram redução "
        "significativa de carga muscular e fadiga quando as tarefas são "
        "adequadamente redistribuídas — Wu et al. (2025) demonstram "
        "redução de carga física e de tempo operacional em tarefas "
        "agrícolas assistidas por cobots. Por fim, Wolffgramm et al. "
        "(2024) mostram que diferentes níveis de latitude decisória "
        "alteram significativamente a percepção do trabalho: ambientes "
        "rígidos reduzem autonomia e motivação, mas autonomia excessiva "
        "gera sobrecarga e variabilidade — o ideal é autonomia calibrada "
        "ao perfil do operador e à complexidade da tarefa.",
    ],
    # §3.4 — compactação leve (Proia/Faccio são retomados em §6.4 e §7.3)
    "3.4 Segurança e ergonomia como variáveis de decisão": [
        "A segurança sempre foi elemento central da robótica industrial, "
        "mas assume novos contornos na robótica colaborativa, em que "
        "humanos e robôs compartilham espaço físico. Tornam-se relevantes "
        "velocidade adaptativa, distância entre operador e robô, zonas "
        "seguras dinâmicas, limitação de força e potência, detecção de "
        "contato, parada segura e previsibilidade de movimento — "
        "requisitos orientados em parte pelas normas ISO 10218 e "
        "ISO/TS 15066.",
        "Proia et al. (2025) propuseram modelo multiobjetivo que integra "
        "tempo de execução, segurança normativa e ergonomia, reforçando "
        "que produtividade isolada não é critério suficiente. Faccio et "
        "al. (2025) mostram que a alocação ótima de tarefas depende de "
        "velocidade do cobot, distância ao operador e risco operacional, "
        "indicando que programação e segurança são dimensões inseparáveis.",
    ],
    # §3.7 — compressão leve (síntese que repete o que já foi dito)
    "3.7 Síntese da revisão": [
        "A literatura confirma que a programação de cobots evoluiu de uma "
        "lógica puramente técnica para uma problemática sociotécnica. O "
        "método ideal depende simultaneamente de: natureza da tarefa, "
        "variabilidade, exigências de segurança, ergonomia, perfil do "
        "operador, confiança, maturidade digital da empresa e metas de "
        "desempenho. Permanecem, todavia, escassos os frameworks capazes "
        "de operacionalizar essa decisão de forma integrada — lacuna que "
        "fundamenta a proposição do Adaptive Human-Centered Framework "
        "for Cobot Programming Selection (AHCF-CPS), apresentado nas "
        "seções seguintes.",
    ],
    # §5.3.1..§5.3.6 — enxugar redundância com §3.3 e §6.x
    "5.3.1 Métodos de programação intuitiva": [
        "Essa categoria reúne estudos que buscam reduzir barreiras técnicas "
        "para programação por operadores não especialistas: Programming by "
        "Demonstration, ensino cinestésico, interfaces gráficas "
        "simplificadas, programação por XR e demonstração multimodal. "
        "Meattini et al. (2025) demonstram que o ensino cinestésico pode "
        "incorporar sEMG e feedback vibrotátil, configurando não só "
        "trajetórias mas também comportamento mecânico do robô, e "
        "estudos com realidade estendida indicam melhora na usabilidade "
        "percebida, embora tarefas espaciais complexas ainda imponham "
        "desafios de precisão. Esses resultados sugerem substituição "
        "gradual da programação rígida baseada em código por métodos de "
        "baixo atrito cognitivo, mais acessíveis ao chão de fábrica.",
    ],
    "5.3.2 Segurança dinâmica e conformidade normativa": [
        "A segunda categoria concentra pesquisas que tratam segurança "
        "como variável central de programação e operação. Proia et al. "
        "(2025) integram simultaneamente tempo de trajetória, requisitos "
        "normativos de segurança e ergonomia humana, mostrando que a "
        "otimização exclusiva de produtividade pode gerar soluções "
        "inadequadas quando desconsidera a exposição humana ao risco. "
        "Faccio et al. (2025) reforçam que a distribuição ótima de "
        "tarefas entre humano e robô depende de distância relativa, "
        "velocidade adaptativa e risco operacional. A segurança, "
        "portanto, não deve ser entendida como camada posterior de "
        "proteção, mas como variável estrutural na escolha do método de "
        "programação, favorecendo modelos adaptativos capazes de ajustar "
        "comportamento em tempo real.",
    ],
    "5.3.3 Ergonomia e saúde ocupacional": [
        "A terceira categoria reúne estudos sobre impactos biomecânicos e "
        "ergonômicos da colaboração humano-robô. Wu et al. (2025) "
        "verificam redução significativa de carga física e de tempo "
        "operacional em tarefas agrícolas assistidas por cobots; "
        "resultados semelhantes aparecem em estudos industriais que "
        "transferem ao robô atividades repetitivas, posturas "
        "desfavoráveis ou manipulação de cargas. Os dados indicam que a "
        "programação de cobots não deve buscar apenas eficiência "
        "temporal, mas redistribuição inteligente de esforço físico.",
    ],
    "5.3.4 Confiança, transparência e usabilidade": [
        "A quarta categoria contempla fatores subjetivos da interação "
        "humano-robô. Bagheri et al. (2022) demonstram que interfaces "
        "transparentes, capazes de explicar por que o cobot escolheu "
        "determinada ação, aumentam significativamente confiança, "
        "satisfação e capacidade de ensinar novas instruções — a "
        "explicabilidade atua como mecanismo funcional de desempenho, "
        "não apenas como atributo comunicacional. Estudos de XR também "
        "apontam maior usabilidade percebida em interfaces visuais e "
        "interativas. Métodos tecnicamente sofisticados, porém opacos "
        "ou pouco intuitivos, tendem a encontrar maior resistência "
        "organizacional.",
    ],
    "5.3.5 Autonomia e latitude decisória": [
        "Wolffgramm et al. (2024) mostram que diferentes configurações de "
        "latitude decisória alteram a percepção do trabalho, a "
        "produtividade e o bem-estar ocupacional: ambientes rígidos "
        "reduzem autonomia percebida, enquanto liberdade irrestrita pode "
        "gerar sobrecarga e inconsistência. O melhor desempenho decorre "
        "de autonomia calibrada ao contexto, não de autonomia máxima "
        "de um dos lados.",
    ],
    "5.3.6 Adaptação às preferências humanas": [
        "Noormohammadi-Asl et al. (2025) mostram que alguns operadores "
        "preferem liderar a tarefa enquanto outros preferem seguir "
        "decisões do sistema; modelos capazes de adaptar-se a essas "
        "preferências mantêm o desempenho e elevam satisfação. O "
        "resultado rompe com a ideia de um operador médio universal e "
        "reforça a necessidade de sistemas personalizados.",
    ],
    # §5.8 — compressão da síntese final (já repete 5.3, 5.4 e 5.5)
    "5.8 Síntese final da análise": [
        "A análise dos dados demonstra que a programação de cobots evoluiu "
        "de uma atividade predominantemente técnica para uma decisão "
        "sociotécnica complexa. O método ideal depende da interação entre "
        "tecnologia disponível, características da tarefa, exigências de "
        "segurança, bem-estar ocupacional, preferências humanas e metas "
        "produtivas — o que fundamenta a proposição de um framework "
        "adaptativo e centrado no humano.",
    ],
    # §6.1 — compactação sem perder citações
    "6.1 Evolução recente da literatura sobre programação de cobots": [
        "Entre 2019 e 2026, a literatura sobre programação de cobots deixou "
        "de se concentrar exclusivamente em controle de trajetória e "
        "segurança básica e passou a incorporar, de forma mais explícita, "
        "dimensões humanas, ergonômicas e adaptativas. El Zaatari et al. "
        "(2019) já apontavam que a programação para tarefas colaborativas "
        "industriais exigia dois requisitos simultâneos: um componente "
        "intuitivo, que permitisse ao operador criar ou modificar programas "
        "com facilidade, e um componente human-aware, capaz de tornar o "
        "robô flexível e adaptável ao parceiro humano; os autores "
        "organizaram o campo em comunicação, otimização e aprendizagem.",
        "Nos estudos mais recentes, observa-se deslocamento do foco da "
        "simples programação para a seleção contextual do método. Tassi, "
        "De Momi e Ajoudani (2022) mostram que a colaboração eficiente "
        "requer controle hierárquico capaz de integrar produtividade com "
        "preferências e conforto humano; Calzavara et al. (2024) defendem "
        "que o desenho de sistemas colaborativos precisa considerar "
        "simultaneamente produtividade, flexibilidade e bem-estar, "
        "introduzindo alocação dinâmica de tarefas baseada em dados "
        "humanos em tempo real. O primeiro resultado geral é, portanto, "
        "uma ampliação do objeto de estudo: de métodos isolados para "
        "ecossistemas de decisão sociotécnica em HRC.",
    ],
    # §6.2 — mantém o quadro e compacta o texto analítico pós-quadro
    "6.2 Taxonomia dos métodos de programação de cobots": [
        "A análise dos artigos permitiu organizar os métodos em seis "
        "classes principais. A primeira é a programação convencional, "
        "ainda presente em ambientes de baixa variabilidade, mas "
        "reconhecidamente limitada em cenários high-mix low-volume, de "
        "reconfiguração frequente ou colaboração dinâmica (EL ZAATARI et "
        "al., 2019). A segunda é a programação parametrizada, na qual o "
        "operador ajusta variáveis de comportamento sem alterar código "
        "em baixo nível: Giberti et al. (2022) descrevem metodologia de "
        "implementação flexível baseada em processo piramidal "
        "parametrizado com refinamento interativo por usuários não "
        "especialistas, e Wolffgramm et al. (2024) mostram operadores "
        "utilizando latitudes decisórias — ajuste de velocidade e "
        "realocação de tarefas — para redesenhar a interdependência "
        "humano-cobot.",
        "A terceira classe é Programming by Demonstration: Arrais et al. "
        "(2021), em coating industrial, mostram que a demonstração "
        "intuitiva transfere know-how de operadores experientes para o "
        "robô; Al-Yacoub et al. (2021) reforçam o valor do aprendizado "
        "por força/torque em tarefas colaborativas de manipulação e "
        "montagem. A quarta classe amplia o PbD para o ensino cinestésico "
        "multimodal — Meattini et al. (2025) mostram que o operador "
        "pode ensinar não só trajetória, mas também compliance e "
        "interação do robô, usando sEMG e feedback vibrotátil, "
        "deslocando a programação do plano geométrico para o plano "
        "relacional e dinâmico.",
        "A quinta classe reúne interfaces multimodais e imersivas: "
        "Angleraud et al. (2021) mostram que comandos gráficos e de fala "
        "coordenam tarefas compartilhadas quando o conjunto de ações é "
        "conhecido e atualizado em tempo de execução; Chan et al. (2022) "
        "mostram que uma interface AR-HMD reduz demanda física, melhora "
        "utilização do robô e reduz tempo de conclusão em tarefas "
        "fisicamente compartilhadas — embora nem sempre seja percebida "
        "como mais confiável; Nguyen et al. (2026) reforçam o potencial "
        "da visão computacional e do reconhecimento contínuo de gestos. "
        "A sexta classe reúne planejamento adaptativo e aprendizado "
        "incremental: Bagheri et al. (2022) mostram que interfaces "
        "transparentes melhoram o ensino de novas instruções por "
        "usuários não especialistas, e Noormohammadi-Asl et al. (2025) "
        "demonstram que o robô pode adaptar alocação e planejamento ao "
        "estilo humano de liderar ou seguir, combinando preferência e "
        "desempenho.",
        "A taxonomia permite entender a programação de cobots como um "
        "continuum de métodos, variando em grau de especialização, "
        "autonomia, transparência e participação do operador.",
    ],
    # §6.4 — compressão (retoma conceitos de §3.4, Proia e Faccio)
    "6.4 Segurança como variável estrutural, não acessória": [
        "A segurança aparece não como requisito externo ao método de "
        "programação, mas como variável interna à própria lógica de "
        "seleção e operação. Saenz et al. (2023) mostram que a segurança "
        "segue sendo barreira relevante à expansão do uso de cobots, "
        "especialmente quando stakeholders têm pouca experiência em "
        "engenharia de segurança — aplicações colaborativas exigem "
        "suporte prático para compreender protocolos, habilidades de "
        "segurança e validação ao longo do ciclo de vida.",
        "Faccio, Granata e Minto (2024) demonstram que a alocação de "
        "tarefas em células colaborativas deve incorporar restrição de "
        "segurança, considerando velocidade variável do cobot e distância "
        "entre recursos, e argumentam que produtividade e segurança não "
        "devem ser tratadas como objetivos isolados. Proia et al. (2025) "
        "aprofundam essa linha com planejamento de trajetória "
        "multiobjetivo capaz de equilibrar tempo de travessia, ergonomia "
        "e Speed and Separation Monitoring, respeitando requisitos ISO: "
        "o melhor desempenho industrial não corresponde ao menor tempo "
        "possível, mas ao melhor compromisso entre eficiência, segurança "
        "e ergonomia. Métodos de programação que ignoram a dinâmica de "
        "segurança tendem, portanto, a ser insuficientes para ambientes "
        "reais de HRC.",
    ],
    # §6.5 — compressão
    "6.5 Ergonomia e bem-estar do operador como critérios crescentes de seleção": [
        "Cresce o número de estudos que integram ergonomia ao desenho, "
        "controle e programação de cobots. Tassi, De Momi e Ajoudani "
        "(2022) propõem um framework de controle hierárquico com "
        "compliance adaptativo, tratando postura e conforto do "
        "trabalhador como variáveis centrais. Calzavara et al. (2024) "
        "desenvolvem estratégia dinâmica de alocação que minimiza "
        "makespan e gasto energético do operador, transferindo tarefas "
        "do humano para o cobot quando há sinais de maior estresse — "
        "uma lógica fortemente alinhada à Indústria 5.0. Em contexto "
        "agrícola, Wu, Wang e Hu (2025) encontram redução significativa "
        "de velocidade e aceleração na coluna e redução do tempo de "
        "execução sob assistência, com redistribuição da ativação "
        "muscular: o ganho ergonômico não é simples eliminação de "
        "esforço, mas reorganização biomecânica da tarefa. A literatura "
        "passa a reconhecer a ergonomia não apenas como benefício "
        "potencial, mas como critério de projeto e de escolha do método.",
    ],
    # §6.6 — compressão
    "6.6 Transparência, confiança e usabilidade influenciam diretamente a eficácia do método": [
        "Fatores subjetivos — especialmente transparência, confiança e "
        "usabilidade — influenciam de forma decisiva o sucesso da "
        "programação colaborativa. Bagheri et al. (2022) demonstram "
        "empiricamente que interfaces transparentes melhoram o desempenho "
        "objetivo ao ensinar novas instruções ao cobot: participantes que "
        "usaram a interface transparente ensinaram com mais acurácia as "
        "restrições faltantes do problema de montagem, com maior "
        "satisfação e confiança. Chan et al. (2022) identificam que a "
        "interface AR reduziu demanda física e tempo de execução e "
        "aumentou a utilização do robô, mas interfaces convencionais "
        "ainda foram percebidas como mais confiáveis em alguns aspectos. "
        "Noormohammadi-Asl et al. (2025) mostram que o robô pode adaptar "
        "seu planejamento às preferências humanas de liderar ou seguir, "
        "sem sacrificar desempenho. Construtos humano-robô são, "
        "portanto, determinantes — não periféricos — para a viabilidade "
        "do método em HRC.",
    ],
    # §6.7 — compressão
    "6.7 Complexidade da tarefa e habilidade do operador como moderadores": [
        "Dois fatores aparecem repetidamente como moderadores da escolha "
        "do método: complexidade da tarefa e habilidade do operador. Em "
        "tarefas simples e repetitivas, métodos estruturados ou "
        "parametrizados tendem a ser suficientes; em tarefas com "
        "múltiplas restrições, mudança frequente de sequência, "
        "variabilidade dimensional ou interação física sensível, a "
        "literatura privilegia PbD, interfaces explicáveis e "
        "planejamento reativo. Wolffgramm et al. (2024) mostram que, à "
        "medida que se amplia a autonomia do operador para reconfigurar "
        "o comportamento do cobot, cresce também a necessidade de "
        "assistência instrumental e suporte — métodos mais sofisticados "
        "não são universalmente vantajosos e sua eficácia depende da "
        "maturidade do usuário. Esse resultado reforça a hipótese "
        "central: não existe um único método ótimo; existe adequação "
        "contingencial entre método, tarefa, operador e ambiente.",
    ],
    # §6.9 — compressão da síntese dos resultados
    "6.9 Síntese dos resultados": [
        "Os principais resultados da revisão podem ser sintetizados em seis "
        "pontos: (i) a literatura recente migrou de uma visão puramente "
        "técnica para uma visão sociotécnica centrada no humano; (ii) os "
        "métodos identificados organizam-se em seis classes — convencional, "
        "parametrizada, Programming by Demonstration, ensino cinestésico "
        "multimodal, interfaces imersivas e planejamento adaptativo; "
        "(iii) segurança e ergonomia aparecem como variáveis estruturais, "
        "não apenas acessórias; (iv) transparência, confiança, usabilidade "
        "e autonomia calibrada afetam diretamente a eficácia do método; "
        "(v) complexidade da tarefa e habilidade do operador moderam a "
        "adequação de cada abordagem; (vi) persiste uma lacuna clara "
        "quanto à existência de um framework integrado de seleção.",
        "Esses resultados sustentam a proposição do Adaptive Human-Centered "
        "Framework for Cobot Programming Selection (AHCF-CPS), apresentado "
        "a seguir como resposta teórica e aplicada às lacunas identificadas.",
    ],
    # §7.2 — compressão do elenco redundante de vantagens/limitações
    "7.2 Não existe método universalmente superior": [
        "Outro ponto central é que não há evidência científica que sustente "
        "a superioridade universal de um único método. Cada abordagem "
        "apresenta vantagens e limitações dependentes do contexto:",
        "Programação offline tradicional — previsibilidade, precisão e "
        "controle fino, mas reprogramação lenta e alta dependência de "
        "especialistas.",
        "Programming by Demonstration — intuitivo e rápido para adaptação, "
        "porém com variabilidade entre demonstrações e sensibilidade à "
        "qualidade do instrutor.",
        "XR e interfaces imersivas — visualização intuitiva, treinamento "
        "remoto e menor demanda física, com desafios em custo e precisão "
        "para tarefas tridimensionais complexas.",
        "Shared control e adaptive AI — alta adaptabilidade em ambientes "
        "dinâmicos, ao custo de maior complexidade sistêmica e exigência "
        "de maturidade digital.",
        "Esse resultado reforça a premissa central do estudo: a melhor "
        "solução depende do alinhamento entre método e contexto, não da "
        "adoção indiscriminada de tendências tecnológicas.",
    ],
    # §7.13 — síntese da discussão, leve ajuste
    "7.13 Síntese da discussão": [
        "A discussão demonstra que a programação de cobots não deve mais "
        "ser tratada como decisão puramente técnica. O estado da arte "
        "indica necessidade de modelos que conciliem desempenho, "
        "segurança, ergonomia, confiança, flexibilidade e adaptação "
        "humana. O AHCF-CPS responde a essa necessidade como modelo "
        "integrado, adaptativo e centrado no humano para orientar a "
        "escolha de métodos de programação de cobots na indústria "
        "contemporânea.",
    ],
    # §8.1 — conclusão compactada
    "8.1 Conclusão": [
        "Este estudo analisou, por meio de uma Revisão Sistemática da "
        "Literatura, os métodos contemporâneos de programação de robôs "
        "colaborativos em HRC industrial e propôs um modelo integrador "
        "para apoio à decisão. A programação de cobots evoluiu de "
        "atividade técnica centrada em especialistas para uma decisão "
        "multidimensional, adaptativa e centrada no ser humano.",
        "Coexistem abordagens distintas — convencional, parametrização, "
        "Programming by Demonstration, ensino cinestésico, interfaces "
        "baseadas em realidade estendida, controle compartilhado e "
        "planejamento adaptativo — sem que nenhuma seja universalmente "
        "superior. A adequação de cada abordagem depende da combinação "
        "entre complexidade e variabilidade da tarefa, exigências de "
        "segurança, condições ergonômicas, habilidade do operador, "
        "confiança e usabilidade percebida, metas de produtividade e "
        "maturidade tecnológica. Isso responde à premissa central: não "
        "existe um único melhor método, mas métodos mais adequados para "
        "contextos específicos.",
        "Fatores humanos — confiança, transparência algorítmica, autonomia "
        "percebida, carga cognitiva e preferências individuais — "
        "consolidam-se como elementos estratégicos da performance "
        "industrial, reforçando a aderência da pesquisa aos princípios "
        "da Indústria 5.0. Com base nessas evidências, propõe-se o "
        "Adaptive Human-Centered Framework for Cobot Programming "
        "Selection (AHCF-CPS), integrando Safety Score, Ergonomic Score, "
        "Human Preference Score, Performance Score e Task Complexity "
        "Score. A principal contribuição teórica está em transformar a "
        "escolha do método em problema multicritério formal; a prática, "
        "em apoiar empresas, integradores e pesquisadores na seleção "
        "contextualizada de estratégias de programação robótica.",
    ],
    # §5.5 — compactação
    "5.5 Lacunas identificadas": [
        "A análise revela quatro lacunas: (a) fragmentação conceitual — "
        "cada estudo privilegia uma dimensão (segurança, ergonomia, IA, "
        "usabilidade ou produtividade), com poucos modelos integradores; "
        "(b) ausência de suporte decisório formal — raramente se trata a "
        "escolha do método como problema multicritério operacionalizável; "
        "(c) baixa transferência industrial — parte dos experimentos "
        "ocorre em laboratório, com tarefas simplificadas e pequenas "
        "amostras; (d) carência de personalização contínua — poucos "
        "sistemas atualizam o método de colaboração conforme mudança do "
        "operador ou do ambiente.",
    ],
    # §5.6 — compactação
    "5.6 Relação com as Questões de Pesquisa": [
        "Para RQ1, os dados confirmam a coexistência de múltiplos métodos — "
        "convencional, PbD, ensino cinestésico, XR, shared control, "
        "planejamento adaptativo e aprendizado incremental. Para RQ2, "
        "evidencia-se influência direta de complexidade da tarefa, "
        "habilidade do operador, segurança, tempo, ergonomia, confiança "
        "e preferência humana. Para RQ3, a combinação dos resultados "
        "mostra base empírica suficiente para propor modelo integrado de "
        "decisão.",
    ],
    # §7.4 — compressão (sobrepõe §3.3)
    "7.4 O papel dos construtos humano-robô": [
        "Fatores subjetivos tradicionalmente considerados “intangíveis” "
        "exercem influência concreta sobre desempenho industrial: "
        "Bagheri et al. (2022) mostram que transparência explicativa "
        "melhora confiança e capacidade de ensino do usuário — sistemas "
        "opacos podem reduzir eficiência mesmo quando tecnicamente "
        "competentes; interfaces intuitivas reduzem curva de aprendizado "
        "e ampliam adoção por não especialistas; métodos excessivamente "
        "complexos transferem esforço físico para esforço mental, "
        "gerando nova forma de sobrecarga; e Noormohammadi-Asl et al. "
        "(2025) demonstram que trabalhadores diferem quanto à "
        "preferência por liderar ou seguir decisões do sistema. Os "
        "fatores humanos não são acessórios emocionais da automação, mas "
        "variáveis produtivas reais.",
    ],
    # §7.6 — compressão
    "7.6 Lacuna consolidada pela revisão": [
        "A principal lacuna consolidada pelos resultados é a ausência, na "
        "literatura recente, de um modelo integrado que responda "
        "simultaneamente: qual método de programação usar, em qual tipo "
        "de tarefa, para qual perfil de operador, sob quais restrições "
        "de segurança, com quais metas de desempenho e considerando "
        "quais impactos ergonômicos. Há bons estudos setoriais e "
        "soluções parciais, mas poucos convertem essas variáveis em um "
        "sistema formal de apoio à decisão. Essa lacuna justifica a "
        "proposição do Adaptive Human-Centered Framework for Cobot "
        "Programming Selection como contribuição teórica e aplicada "
        "deste estudo.",
    ],
    # §7.7 — incorpora o Quadro 3 em texto corrido
    "7.7 Proposição do Adaptive Human-Centered Framework for Cobot Programming Selection": [
        "Propõe-se que a seleção do método seja tratada por um framework "
        "multicritério composto por cinco dimensões centrais, sintetizando "
        "os fatores recorrentes encontrados nos artigos analisados: "
        "Safety Score, para avaliar risco operacional e requisitos "
        "normativos; Ergonomic Score, para medir esforço físico, postura "
        "e fadiga; Human Preference Score, para considerar confiança, "
        "experiência e estilo de trabalho; Performance Score, para "
        "produtividade, qualidade e tempo de ciclo; e Task Complexity "
        "Score, para variabilidade, precisão e necessidade decisória.",
    ],
    # §7.5 — compactação
    "7.5 Implicações para pequenas e médias empresas": [
        "A discussão revela implicações importantes para PMEs industriais. "
        "Enquanto grandes corporações absorvem parte da complexidade com "
        "equipes especializadas, as PMEs operam com recursos limitados, "
        "equipes enxutas, alta pressão por retorno rápido e necessidade "
        "de flexibilidade produtiva. Métodos excessivamente complexos ou "
        "dependentes de especialistas externos tendem a reduzir "
        "viabilidade econômica, enquanto abordagens intuitivas, "
        "parametrizadas ou híbridas aceleram adoção e reduzem barreiras "
        "tecnológicas. Por isso, frameworks de decisão tornam-se ainda "
        "mais relevantes para PMEs: ajudam a evitar investimentos "
        "inadequados.",
    ],
    # §5.2 — compactação
    "5.2 Caracterização geral da amostra": [
        "Os 21 artigos selecionados revelam crescimento consistente da "
        "produção científica sobre cobots entre 2019 e 2026, "
        "especialmente a partir de 2022, período em que temas como "
        "Indústria 5.0, ergonomia digital, explicabilidade algorítmica e "
        "interfaces intuitivas ganharam destaque. Os estudos concentram-se "
        "em periódicos e conferências das áreas de robótica aplicada, "
        "automação industrial, engenharia de produção, interação "
        "humano-computador, inteligência artificial e ergonomia "
        "ocupacional, com predominância de pesquisas oriundas de Europa, "
        "Ásia e América do Norte.",
    ],
    # §7.3 — compressão
    "7.3 Segurança e ergonomia como critérios de primeira ordem": [
        "A literatura analisada mostra consistentemente que segurança e "
        "ergonomia não devem ser tratadas como critérios secundários nem "
        "como verificações posteriores à implantação, mas como variáveis "
        "de primeira ordem na decisão sobre programação de cobots. "
        "Proia et al. (2025) demonstram que trajetórias ótimas em tempo "
        "podem ser inadequadas quando expõem operadores a maior risco ou "
        "posturas desfavoráveis, e Faccio et al. (2024) mostram que a "
        "distribuição de tarefas em células colaborativas depende de "
        "velocidade, distância e risco operacional. Na prática, muitas "
        "empresas avaliam a automação colaborativa apenas pelo ROI "
        "direto, ignorando custos ocultos como acidentes, fadiga, "
        "absenteísmo, retrabalho e baixa aceitação do operador. A "
        "programação de cobots deve, portanto, incorporar desde o início "
        "métricas de segurança dinâmica e de ergonomia ocupacional.",
    ],
    # §7.1 — compressão
    "7.1 Superação da visão tradicional de programação robótica": [
        "Uma primeira implicação é a necessidade de superar a lógica "
        "tradicional de programação robótica baseada em especialistas e "
        "rotinas rígidas. Embora a programação convencional permaneça "
        "útil em processos estáveis e repetitivos, sua adequação cai com "
        "o aumento da variabilidade de produtos, da customização de "
        "lotes, de mudanças frequentes de layout, da necessidade de "
        "reconfiguração rápida e da participação ativa do operador. "
        "El Zaatari et al. (2019) já antecipavam que a colaboração "
        "industrial exige métodos simultaneamente intuitivos e "
        "human-aware; a literatura posterior confirma o diagnóstico ao "
        "expandir PbD, interfaces imersivas, controle compartilhado e "
        "aprendizagem incremental. Programar cobots deixa, assim, de "
        "significar apenas escrever instruções e passa a significar "
        "desenhar uma relação operacional entre humano e robô.",
    ],
    # §8.3 consolidado — subseções viram um único parágrafo denso
    "8.3 Estudos Futuros": [
        "Os resultados abrem diversas oportunidades de aprofundamento. "
        "A principal agenda é testar o AHCF-CPS em ambientes industriais "
        "reais — automotivo, eletroeletrônico, logística, farmacêutico, "
        "agroindustrial e metalmecânico — comparando suas recomendações "
        "com decisões tradicionais de implantação. O modelo multicritério "
        "pode ser expandido com técnicas como AHP, TOPSIS, Fuzzy Logic, "
        "Bayesian Decision Models e Reinforcement Learning adaptativo, "
        "permitindo recomendações automáticas a partir de dados "
        "históricos e sensores em tempo real. Outras frentes incluem "
        "integração a gêmeos digitais industriais — para simular métodos "
        "antes da implantação física — personalização contínua — "
        "ajustando o método conforme mudanças no perfil do operador, "
        "fadiga e aprendizado acumulado — e extensão para múltiplos "
        "cobots, equipes híbridas, logística interna autônoma e "
        "colaboração entre robôs móveis e manipuladores. Alinhado à "
        "agenda da Indústria 5.0, o framework pode ainda incorporar "
        "indicadores ESG e de sustentabilidade, como consumo energético, "
        "redução de acidentes e impacto ocupacional de longo prazo.",
    ],
    # §4.9 + §4.10 + §4.11 + §4.13 — consolidados em prosa
    "4.9 Extração de dados": [
        "Para cada artigo selecionado, elaborou-se planilha padronizada "
        "contendo autores, ano, periódico ou conferência, país/instituição, "
        "setor de aplicação, método de programação estudado, nível de "
        "autonomia, presença de fatores humanos, variáveis analisadas, "
        "principais resultados, limitações e contribuição para o framework "
        "proposto — padronização que permitiu comparabilidade entre "
        "estudos heterogêneos.",
    ],
    "4.10 Estratégia de análise": [
        "A análise combinou análise bibliométrica descritiva — identificando "
        "evolução temporal, periódicos predominantes, palavras-chave "
        "recorrentes e clusters temáticos — com análise de conteúdo "
        "temática que classificou os estudos em segurança e conformidade "
        "normativa, ergonomia e saúde ocupacional, métodos de "
        "programação intuitiva, confiança e transparência, autonomia "
        "adaptativa, desempenho operacional e critérios de decisão "
        "multicritério. Essa etapa foi fundamental para a construção da "
        "taxonomia e do framework final.",
    ],
    "4.11 Confiabilidade metodológica": [
        "Para aumentar a robustez analítica, adotaram-se leitura integral "
        "dos artigos centrais, comparação cruzada entre estudos "
        "convergentes e divergentes, priorização de periódicos indexados "
        "de alta qualidade, exclusão de evidências frágeis ou não "
        "replicáveis, e triangulação entre resultados quantitativos e "
        "qualitativos.",
    ],
    "4.13 Síntese do método": [
        "A Revisão Sistemática da Literatura permitiu mapear o estado da "
        "arte sobre programação de cobots em HRC industrial, identificar "
        "lacunas científicas e reunir base teórica suficiente para propor "
        "o Adaptive Human-Centered Framework for Cobot Programming "
        "Selection (AHCF-CPS). As seções seguintes apresentam os "
        "resultados consolidados e a estrutura formal do framework.",
    ],
    # §4.4 Estratégia de busca — compactação do bloco de keywords
    "4.4 Estratégia de busca": [
        "A construção das strings de busca considerou sinônimos e termos "
        "amplamente utilizados na literatura internacional, combinando "
        "três grupos conceituais: Tecnologia (cobot, collaborative "
        "robot, human-robot collaboration, human-robot interaction); "
        "Programação (programming, robot programming, programming by "
        "demonstration, kinesthetic teaching, teach pendant, shared "
        "control, adaptive planning); Contexto industrial (industry, "
        "manufacturing, production system, industrial environment). A "
        "expressão principal — (“cobot” OR “collaborative robot” OR "
        "“human-robot collaboration”) AND (“programming” OR "
        "“robot programming” OR “programming by demonstration” OR "
        "“kinesthetic teaching” OR “shared control”) AND (“industry” OR "
        "“manufacturing” OR “industrial environment”) — foi adaptada "
        "conforme os operadores específicos de cada base.",
    ],
    # §3.6 Lacunas identificadas — compactação (listas extensas repetem §5.5)
    "3.6 Lacunas identificadas na literatura": [
        "Apesar dos avanços, a revisão evidencia lacunas importantes: "
        "(a) fragmentação temática — estudos focam isoladamente em "
        "segurança, ergonomia, interfaces, IA, produtividade ou "
        "experiência do usuário, com poucos modelos integradores; "
        "(b) ausência de frameworks decisórios que tratem a seleção do "
        "método como problema formal de decisão multicritério; "
        "(c) baixa generalização industrial, pois muitos estudos são "
        "laboratoriais, com tarefas simplificadas e amostras pequenas; "
        "(d) subexploração de diferenças individuais — experiência, "
        "idade, perfil cognitivo e preferências pessoais; e "
        "(e) carência de modelos dinâmicos, já que grande parte da "
        "literatura assume decisões estáticas enquanto ambientes reais "
        "mudam continuamente.",
    ],
    # §8.4 — consideração final bastante enxugada
    "8.4 Consideração Final": [
        "A próxima geração da automação industrial será definida pela "
        "capacidade de integrar tecnologia e trabalho humano de forma "
        "inteligente, segura e sustentável. Nesse contexto, selecionar "
        "corretamente como programar um cobot é tão importante quanto "
        "decidir se utilizá-lo. O AHCF-CPS representa um passo inicial "
        "nessa direção, propondo automação colaborativa eficiente, "
        "adaptativa, inclusiva e centrada nas pessoas.",
    ],
}

# §7.9 é um caso à parte: texto revisado + fórmula OMML
SECTION_79_INTRO = [
    "O escore global de adequação do método j ∈ M, para um cenário "
    "descrito pelas cinco dimensões consolidadas em §5.7 (Safety, "
    "Ergonomic, Human Preference, Performance e Task Complexity), "
    "é dado pela formulação:",
]
SECTION_79_LEGEND_TITLE = "Onde, com domínios explícitos:"
SECTION_79_LEGEND_ITEMS = [
    "X_i ∈ [0, 1] — fatores contextuais; H_k ∈ [0, 1] — construtos "
    "humano-robô. Ambos derivam das cinco dimensões consolidadas em "
    "§5.7, obtidas a partir de escalas Likert 1–5 por "
    "norm(ℓ) = (ℓ − 1)/4.",
    "a_{j,·} ∈ [0, 1] — afinidade do método j com a variável "
    "correspondente, determinada por tabela fundamentada na Revisão "
    "Sistemática (§6).",
    "α_i, β_k, γ ∈ ℝ_{≥0} — pesos ajustáveis, calibráveis por setor.",
    "c_j ∈ [0, 1] — custo relativo do método j.",
    "χ̃ ∈ [0, 1] — restrição orçamentária do cenário.",
]
SECTION_79_OUTRO = [
    "A formulação é diretamente compatível com Analytic Hierarchy "
    "Process — via α_i, β_k obtidos da média geométrica de uma "
    "matriz pareada entre as dimensões — com Technique for Order "
    "Preference by Similarity to Ideal Solution — via solução ideal "
    "sobre a matriz [a_{j,·} X, a_{j,·} H] — com Fuzzy Logic "
    "(variáveis linguísticas sobre X e H) e com aprendizado de "
    "máquina (vetor de entrada determinístico). Em relação ao "
    "rascunho inicial desta seção, as alterações fundamentais são "
    "cinco: (R1) introdução da afinidade a_{j,·}, para que A_j "
    "dependa de j — condição lógica mínima de ranking em decisão "
    "multicritério; (R2) fechamento do domínio de X_i e H_k em "
    "[0, 1], com regra explícita de normalização Likert — condição "
    "de reprodutibilidade; (R3) consolidação da partição X/H nas "
    "cinco dimensões do §5.7 — condição de coerência interna entre "
    "seções; (R4) tratamento do custo como termo de penalidade "
    "(sinal negativo) — condição de consistência com a noção de "
    "restrição orçamentária; (R5) formalização do vínculo "
    "método × variáveis via tabela de afinidades sobre os seis "
    "métodos da taxonomia (§6.2) — condição para que a nota sobre "
    "implementação via AHP, TOPSIS, Fuzzy Logic ou aprendizado de "
    "máquina passe a ter objeto computável.",
]

# ---------------------------------------------------------------------------
# Parser: quebra o .txt em seções
# ---------------------------------------------------------------------------

SECTION_HEADER_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+([A-ZÁÉÍÓÚÃÕÂÊÎÔÛÇ].*)$")


def normalize_header(line: str) -> str | None:
    """Devolve '1. Resumo' → '1. Resumo', '3.1 Evolução...' → '3.1 Evolução...'.

    Retorna None se a linha não parece cabeçalho.
    """
    m = SECTION_HEADER_RE.match(line.strip())
    if not m:
        return None
    number = m.group(1)
    title = m.group(2).strip()
    # cabeçalhos costumam ser curtos
    if len(title) > 110:
        return None
    if title.endswith("."):
        return None  # provavelmente é frase, não título
    return f"{number} {title}" if "." in number else f"{number}. {title}"


def load_sections() -> List[Tuple[str, List[str]]]:
    """Retorna lista de (header, paragraphs)."""
    text = SOURCE.read_text(encoding="utf-8")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    sections: List[Tuple[str, List[str]]] = []
    current_header = "<preamble>"
    current_body: List[str] = []
    for p in paragraphs:
        header = normalize_header(p)
        if header:
            sections.append((current_header, current_body))
            current_header = header
            current_body = []
        else:
            current_body.append(p)
    sections.append((current_header, current_body))
    return sections


# ---------------------------------------------------------------------------
# OMML helpers para a fórmula §7.9
# ---------------------------------------------------------------------------

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def m(tag: str) -> str:
    return f"{{{M_NS}}}{tag}"


def make_math_run(text: str, *, sty: str | None = None) -> OxmlElement:
    r = OxmlElement("m:r")
    if sty:
        rpr = OxmlElement("m:rPr")
        sty_el = OxmlElement("m:sty")
        sty_el.set(qn("m:val"), sty)
        rpr.append(sty_el)
        r.append(rpr)
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def make_sub(base: str, sub: str) -> OxmlElement:
    s = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(make_math_run(base, sty="i"))
    s.append(e)
    sub_el = OxmlElement("m:sub")
    sub_el.append(make_math_run(sub, sty="i"))
    s.append(sub_el)
    return s


def make_sup(base: str, sup: str) -> OxmlElement:
    s = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    e.append(make_math_run(base, sty="i"))
    s.append(e)
    sup_el = OxmlElement("m:sup")
    sup_el.append(make_math_run(sup, sty="p"))
    s.append(sup_el)
    return s


def make_nary(lower: str, body_nodes: List[OxmlElement]) -> OxmlElement:
    """Σ com limites inferior (e sem superior visível)."""
    nary = OxmlElement("m:nary")
    pr = OxmlElement("m:naryPr")
    chr_el = OxmlElement("m:chr")
    chr_el.set(qn("m:val"), "∑")
    pr.append(chr_el)
    limloc = OxmlElement("m:limLoc")
    limloc.set(qn("m:val"), "undOvr")
    pr.append(limloc)
    subhide = OxmlElement("m:subHide")
    subhide.set(qn("m:val"), "0")
    pr.append(subhide)
    suphide = OxmlElement("m:supHide")
    suphide.set(qn("m:val"), "1")
    pr.append(suphide)
    nary.append(pr)
    sub_el = OxmlElement("m:sub")
    sub_el.append(make_math_run(lower, sty="i"))
    nary.append(sub_el)
    sup_el = OxmlElement("m:sup")
    nary.append(sup_el)
    e = OxmlElement("m:e")
    for node in body_nodes:
        e.append(node)
    nary.append(e)
    return nary


def formula_paragraph(doc: Document) -> None:
    """Insere A_j = Σ_i α_i a_{j,i} X_i + Σ_k β_k a_{j,k} H_k - γ c_j χ̃  ;  m* = argmax_j A_j."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    oMathPara = OxmlElement("m:oMathPara")
    oMath = OxmlElement("m:oMath")

    # A_j =
    oMath.append(make_sub("A", "j"))
    oMath.append(make_math_run(" = ", sty="p"))

    # Σ_i α_i a_{j,i} X_i
    body1: List[OxmlElement] = []
    body1.append(make_sub("α", "i"))
    body1.append(make_math_run(" ", sty="p"))
    body1.append(make_sub("a", "j,i"))
    body1.append(make_math_run(" ", sty="p"))
    body1.append(make_sub("X", "i"))
    oMath.append(make_nary("i", body1))

    oMath.append(make_math_run(" + ", sty="p"))

    # Σ_k β_k a_{j,k} H_k
    body2: List[OxmlElement] = []
    body2.append(make_sub("β", "k"))
    body2.append(make_math_run(" ", sty="p"))
    body2.append(make_sub("a", "j,k"))
    body2.append(make_math_run(" ", sty="p"))
    body2.append(make_sub("H", "k"))
    oMath.append(make_nary("k", body2))

    # - γ c_j χ̃
    oMath.append(make_math_run(" − ", sty="p"))
    oMath.append(make_math_run("γ", sty="i"))
    oMath.append(make_math_run(" ", sty="p"))
    oMath.append(make_sub("c", "j"))
    oMath.append(make_math_run(" ", sty="p"))
    oMath.append(make_math_run("χ̃", sty="i"))

    oMathPara.append(oMath)
    p._p.append(oMathPara)

    # segunda linha: m* = argmax_j A_j
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    oMathPara2 = OxmlElement("m:oMathPara")
    oMath2 = OxmlElement("m:oMath")
    oMath2.append(make_sup("m", "*"))
    oMath2.append(make_math_run(" = arg ", sty="p"))
    oMath2.append(make_sub("max", "j ∈ M"))
    oMath2.append(make_math_run(" ", sty="p"))
    oMath2.append(make_sub("A", "j"))
    oMathPara2.append(oMath2)
    p2._p.append(oMathPara2)


# ---------------------------------------------------------------------------
# Render helpers
# ---------------------------------------------------------------------------


def set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), "Arial")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"


def add_author_block(doc: Document, lines: List[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.name = "Arial"


def add_section_header(doc: Document, header: str) -> None:
    # Escolhe nível pelo número de pontos em "x.y.z"
    num = header.split(" ", 1)[0]
    level = num.count(".") + 1
    if level == 1:
        size = Pt(13)
    elif level == 2:
        size = Pt(12)
    else:
        size = Pt(11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(header)
    r.bold = True
    r.font.size = size
    r.font.name = "Arial"


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.first_line_indent = Pt(-14)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------------------------------------------------------------------------
# Heurística: o parser inclui muitos bullet-points como parágrafos curtos.
# Detectamos parágrafos curtos terminados em `;` ou iniciados por "a)", "•"
# e os renderizamos como bullets.
# ---------------------------------------------------------------------------

BULLET_HINTS_RE = re.compile(r"^[a-z]\)\s|^•\s|^–\s|^-\s")


def looks_like_bullet(text: str) -> bool:
    stripped = text.strip()
    if BULLET_HINTS_RE.match(stripped):
        return True
    if stripped.endswith(";") and len(stripped) < 120:
        return True
    if len(stripped) < 60 and not stripped.endswith("."):
        return True
    return False


# ---------------------------------------------------------------------------
# Render do §7.9
# ---------------------------------------------------------------------------


def render_section_79(doc: Document) -> None:
    add_section_header(doc, "7.9 Formulação matemática")
    for para in SECTION_79_INTRO:
        add_body(doc, para)
    formula_paragraph(doc)
    add_body(doc, SECTION_79_LEGEND_TITLE)
    for item in SECTION_79_LEGEND_ITEMS:
        add_bullet(doc, item)
    for para in SECTION_79_OUTRO:
        add_body(doc, para)


# ---------------------------------------------------------------------------
# Seções especiais que precisam de tratamento (tabelas, capas, etc.)
# ---------------------------------------------------------------------------

SKIP_HEADERS = {
    "4.8 Fluxograma PRISMA",
    # Subseções consolidadas em §8.3 para compactação
    "8.3.1 Validação empírica do framework",
    "8.3.2 Modelagem quantitativa avançada",
    "8.3.3 Integração com Digital Twins",
    "8.3.4 Personalização contínua",
    "8.3.5 Expansão para múltiplos robôs",
    "8.3.6 Indicadores ESG e sustentabilidade",
    # Subseções consolidadas em §3.3
    "3.3.1 Confiança",
    "3.3.2 Usabilidade",
    "3.3.3 Carga cognitiva",
    "3.3.4 Ergonomia",
    "3.3.5 Autonomia e percepção do trabalho",
}


def render_preamble(doc: Document, paragraphs: List[str]) -> None:
    """Capa + Abstract + Keywords. §1/§2/§3 caem no parser como seções."""
    if not paragraphs:
        return
    title = paragraphs[0]
    add_title(doc, title)
    if len(paragraphs) >= 3:
        add_author_block(doc, [paragraphs[1], paragraphs[2]])
    for p in paragraphs[3:]:
        stripped = p.strip()
        if stripped.startswith("Abstract:"):
            h = doc.add_paragraph()
            r = h.add_run("Abstract")
            r.bold = True
            r.font.name = "Arial"
            add_body(doc, stripped.removeprefix("Abstract:").strip())
        elif stripped.startswith("Keywords:"):
            h = doc.add_paragraph()
            r = h.add_run("Keywords")
            r.bold = True
            r.font.name = "Arial"
            add_body(doc, stripped.removeprefix("Keywords:").strip())
        else:
            add_body(doc, stripped)


def render_section(doc: Document, header: str, body: List[str]) -> None:
    if header == "4.8 Fluxograma PRISMA":
        add_section_header(doc, header)
        add_body(
            doc,
            "(O fluxograma PRISMA acompanha o material suplementar do artigo.)",
        )
        return
    if header in SKIP_HEADERS:
        return

    if header == "7.9 Formulação matemática inicial":
        render_section_79(doc)
        return

    # Substituição cirúrgica por versão mais concisa
    if header in REPLACEMENTS:
        add_section_header(doc, header)
        for para in REPLACEMENTS[header]:
            add_body(doc, para)
        return

    # Caso padrão
    add_section_header(doc, header)
    for para in body:
        if para.strip() in STRIP_PARAGRAPHS:
            continue
        if looks_like_bullet(para):
            add_bullet(doc, para)
        else:
            add_body(doc, para)


def render_references(doc: Document, refs: List[str]) -> None:
    add_section_header(doc, "Referências")
    for ref in refs:
        if not ref.strip():
            continue
        add_reference(doc, ref.strip())


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    sections = load_sections()

    doc = Document()
    set_base_style(doc)
    # Margens padrão OK; A4
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(72)
        section.left_margin = section.right_margin = Pt(72)

    # Preâmbulo (título, abstract, §1, §2) + header de §3
    pre_header, pre_body = sections[0]
    assert pre_header == "<preamble>"
    render_preamble(doc, pre_body)

    # Demais seções — separa referências
    ref_index = None
    for i, (h, _b) in enumerate(sections):
        if h.startswith("Referências"):
            ref_index = i
            break

    main_sections = sections[1:ref_index] if ref_index else sections[1:]
    for h, b in main_sections:
        render_section(doc, h, b)

    # Seção de Referências é detectada como "<preamble>" no parser quando o
    # rótulo não bate o padrão numerado. Tratamos à parte:
    # O arquivo original coloca "Referências (seleção)" como título com espaço
    # e os itens seguem como parágrafos. Vamos localizar manualmente.
    text = SOURCE.read_text(encoding="utf-8")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    try:
        idx = next(
            i
            for i, p in enumerate(paragraphs)
            if p.startswith("Referências")
        )
        refs = paragraphs[idx + 1 :]
        render_references(doc, refs)
    except StopIteration:
        pass

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))

    # Contagem de palavras no resultado
    new_text = Path(str(OUT)).read_bytes()
    _ = new_text  # não contamos binário; usamos outra rota abaixo
    # Contagem textual a partir do que foi gerado
    count = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Gerado: {OUT}")
    print(f"Palavras (parágrafos do docx): {count}")


if __name__ == "__main__":
    main()
